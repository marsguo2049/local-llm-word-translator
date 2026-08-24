#!/usr/bin/env python3
"""Local, resumable DOCX translation through LM Studio's native REST API.

The original document is never modified.  Translation data is stored separately
and can be rendered into a Chinese reading copy and an interleaved bilingual copy.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import os
import re
import shutil
import sys
import tempfile
import time
import zipfile
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET

import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
CONFIG_PATH = ROOT / "config.json"
INPUT_DIR = ROOT / "input"
OUTPUT_DIR = ROOT / "output"
PROGRESS_DIR = ROOT / "progress"
GLOSSARY_DIR = ROOT / "glossary"
LOG_DIR = ROOT / "logs"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS}
W_ID = f"{{{W_NS}}}id"

WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ]+(?:[-'’][A-Za-zÀ-ÖØ-öø-ÿ]+)*")
BLOCK_RE = re.compile(
    r"(?ms)^\s*\[(?P<id>[A-Za-z0-9_-]+)\]\s*(?P<text>.*?)"
    r"(?=^\s*\[[A-Za-z0-9_-]+\]\s*|\Z)"
)

STATIC_TRANSLATIONS = {
    "abstract": "摘要",
    "contents": "目录",
    "list of illustrations": "插图目录",
    "introduction": "引言",
    "conclusion": "结论",
    "bibliography": "参考文献",
    "primary sources": "一手文献",
    "secondary sources": "二手文献",
    "illustrations": "插图",
}

SYSTEM_PROMPT = """你是一名严谨的英译中学术翻译员，专业领域为艺术史、视觉文化、宗教艺术和博物馆研究。

规则：
1. 忠实、完整地翻译，不概括、不扩写、不评论，不增加原文没有的背景知识。
2. 保留年份、图号、脚注编号、页码、网址、括号内容和文献引用。
3. 人名、作品名、机构名和术语优先使用通行译名；没有可靠译名时保留原文。
4. 首次出现的重要专名可用“中文（原文）”；作品名称可用《中文名》（Original Title）。
5. 对脚注中的解释性文字进行翻译，但书名、期刊名、出版信息、URL 和引文数据保持原样。
6. 保持学术论文语气。只输出指定编号及译文，不输出解释、摘要或翻译过程。
7. 每个输入编号必须且只能出现一次，不得合并、遗漏或改写编号。"""

TERM_PROMPT = """你是艺术史论文术语编辑。请从文本中提取影响全文一致性的候选术语，并建议简体中文译法。

只提取：人名、艺术作品名、地点与机构、宗教/礼仪术语、艺术史概念、绘画材料与技法、反复出现的关键短语。
不要提取普通英文单词，不要解释正文。没有可靠中文译名时，target 保留原文。

只返回 JSON 数组，不要 Markdown。每项格式：
{"source":"英文原词","target":"建议中文","category":"person|artwork|place|institution|religion|technique|concept|other"}
每批最多 20 项。"""

GLOSSARY_AUDIT_PROMPT = """你是一名谨慎的艺术史英汉术语审校员。你会收到现有术语、原文上下文和相关词条。

任务不是重新翻译论文，而是逐条核对现有译名、识别同一实体的写法变体，并指出必须由人确认的歧义。

规则：
1. 必须为每个输入 source 返回且只返回一项，source 必须原样复制。
2. 人名、地名、修道院、教堂、艺术作品、天主教礼仪及绘画技法优先采用可靠的艺术史/宗教史通行译法。
3. 如果可靠中文译名不确定，recommended_target 保留 current_target，decision 设为 review；不要自信地猜译。
4. canonical_source 用于标记同一实体或概念的规范英文名。不同写法属于同一对象时使用相同 canonical_source，但不要删除原始 source。
5. keep 表示当前译名可直接使用；replace 只用于当前译名明显错误且替代译名把握很高；review 表示存在语境、流派、圣名译法或学界惯例歧义。
6. confidence 是 0 到 1 的数字。reason 使用不超过 35 个中文字符说明关键原因。
7. 不要把作者姓氏自动扩展成全名；短写和全名可以属于同一 canonical_source，但各自译法应符合原文写法。
8. 只返回 JSON 数组，不要 Markdown 或额外说明。

每项格式：
{"source":"原样复制","current_target":"当前译名","recommended_target":"建议译名","canonical_source":"规范英文名","category":"person|artwork|place|institution|religion|technique|concept|other","decision":"keep|replace|review","confidence":0.0,"reason":"简短理由"}
"""


class TranslatorError(RuntimeError):
    pass


class PegNativeFormatError(TranslatorError):
    """LM Studio generated text but rejected it in its peg-native parser."""

    pass


def ensure_dirs() -> None:
    for path in (INPUT_DIR, OUTPUT_DIR, PROGRESS_DIR, GLOSSARY_DIR, LOG_DIR):
        path.mkdir(parents=True, exist_ok=True)


def load_config() -> dict[str, Any]:
    try:
        data = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise TranslatorError(f"缺少配置文件：{CONFIG_PATH}") from exc
    required = ("server_url", "model", "chunk_words", "max_output_tokens")
    missing = [key for key in required if key not in data]
    if missing:
        raise TranslatorError(f"config.json 缺少字段：{', '.join(missing)}")
    return data


def atomic_write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(prefix=path.name, suffix=".tmp", dir=path.parent)
    try:
        with os.fdopen(fd, "w", encoding="utf-8", newline="\n") as handle:
            json.dump(data, handle, ensure_ascii=False, indent=2)
        os.replace(temp_name, path)
    finally:
        if os.path.exists(temp_name):
            os.unlink(temp_name)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def word_count(text: str) -> int:
    return len(WORD_RE.findall(text))


def normalize(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).casefold()


def term_occurrences(source: str, text: str) -> list[re.Match[str]]:
    """Match whole terms; preserve case for title/proper-name entries."""
    first_cased = next((char for char in source if char.isalpha()), "")
    flags = 0 if first_cased.isupper() else re.I
    return list(re.finditer(rf"(?<!\w){re.escape(source)}(?!\w)", text, flags=flags))


def resolve_docx(value: str) -> Path:
    candidate = Path(value)
    if not candidate.is_absolute():
        candidate = ROOT / candidate
    candidate = candidate.resolve()
    if not candidate.exists():
        raise TranslatorError(f"找不到 Word 文件：{candidate}")
    if candidate.suffix.casefold() != ".docx":
        raise TranslatorError("输入文件必须是 .docx，不能是旧版 .doc。")
    return candidate


def state_dir_for(docx_path: Path) -> Path:
    return PROGRESS_DIR / docx_path.stem


def extract_notes(docx_path: Path, kind: str) -> list[dict[str, Any]]:
    part_name = f"word/{kind}.xml"
    tag = "footnote" if kind == "footnotes" else "endnote"
    with zipfile.ZipFile(docx_path) as archive:
        if part_name not in archive.namelist():
            return []
        root = ET.fromstring(archive.read(part_name))

    result: list[dict[str, Any]] = []
    for node in root.findall(f"w:{tag}", NS):
        note_id = int(node.attrib.get(W_ID, "-1"))
        if note_id < 0:
            continue
        paragraphs: list[str] = []
        for paragraph in node.findall(".//w:p", NS):
            text = "".join(t.text or "" for t in paragraph.findall(".//w:t", NS)).strip()
            if text:
                paragraphs.append(text)
        text = "\n".join(paragraphs).strip()
        if text:
            result.append({"note_id": note_id, "source": text})
    return result


def locate_sections(paragraphs: list[str]) -> dict[str, int | None]:
    normalized = [normalize(text) for text in paragraphs]

    def first_exact(value: str, start: int = 0) -> int | None:
        wanted = value.casefold()
        for index in range(start, len(normalized)):
            if normalized[index] == wanted:
                return index
        return None

    contents = first_exact("contents")
    abstract = first_exact("abstract", (contents or 0) + 1)
    bibliography = first_exact("bibliography", (abstract or 0) + 1)
    illustrations = first_exact("illustrations", (bibliography or 0) + 1)
    return {
        "contents": contents,
        "abstract": abstract,
        "bibliography": bibliography,
        "illustrations": illustrations,
    }


def should_translate_paragraph(index: int, text: str, sections: dict[str, int | None]) -> bool:
    if not text.strip():
        return False
    normalized = normalize(text)
    if normalized in STATIC_TRANSLATIONS:
        return True

    contents = sections["contents"]
    abstract = sections["abstract"]
    bibliography = sections["bibliography"]
    illustrations = sections["illustrations"]

    # Translate meaningful title-page text, but not isolated identifiers.
    if contents is not None and index < contents:
        return word_count(text) >= 3
    # Translate the displayed table of contents; Word page numbers must be refreshed later.
    if contents is not None and abstract is not None and contents <= index < abstract:
        return word_count(text) >= 1
    # Translate abstract, illustration list, introduction, chapters and conclusion.
    if abstract is not None and bibliography is not None and abstract <= index <= bibliography:
        return True
    # Leave bibliography entries untouched, except the static section headings above.
    if bibliography is not None and illustrations is not None and bibliography < index < illustrations:
        return False
    # Translate the final illustrations heading and captions while retaining images.
    if illustrations is not None and index >= illustrations:
        return True
    return False


def build_manifest(docx_path: Path, config: dict[str, Any]) -> dict[str, Any]:
    document = Document(docx_path)
    paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    sections = locate_sections(paragraph_texts)
    items: list[dict[str, Any]] = []

    for index, paragraph in enumerate(document.paragraphs):
        source = paragraph.text.strip()
        if not should_translate_paragraph(index, source, sections):
            continue
        item_id = f"P{index:05d}"
        items.append(
            {
                "id": item_id,
                "kind": "paragraph",
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "source": source,
                "words": word_count(source),
            }
        )

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    source = paragraph.text.strip()
                    if word_count(source) < 1:
                        continue
                    item_id = f"T{table_index:03d}R{row_index:03d}C{cell_index:03d}P{paragraph_index:03d}"
                    items.append(
                        {
                            "id": item_id,
                            "kind": "table_cell",
                            "table_index": table_index,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_index": paragraph_index,
                            "source": source,
                            "words": word_count(source),
                        }
                    )

    if config.get("translate_footnotes", True):
        for note in extract_notes(docx_path, "footnotes"):
            items.append(
                {
                    "id": f"FN{note['note_id']:05d}",
                    "kind": "footnote",
                    "note_id": note["note_id"],
                    "source": note["source"],
                    "words": word_count(note["source"]),
                }
            )

    if config.get("translate_endnotes", True):
        for note in extract_notes(docx_path, "endnotes"):
            items.append(
                {
                    "id": f"EN{note['note_id']:05d}",
                    "kind": "endnote",
                    "note_id": note["note_id"],
                    "source": note["source"],
                    "words": word_count(note["source"]),
                }
            )

    manifest = {
        "schema_version": 1,
        "source_path": str(docx_path),
        "source_name": docx_path.name,
        "source_sha256": sha256_file(docx_path),
        "sections": sections,
        "document": {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "inline_shapes": len(document.inline_shapes),
            "sections": len(document.sections),
        },
        "items": items,
    }
    return manifest


def load_or_create_manifest(docx_path: Path, config: dict[str, Any]) -> dict[str, Any]:
    state_dir = state_dir_for(docx_path)
    manifest_path = state_dir / "manifest.json"
    current_hash = sha256_file(docx_path)
    if manifest_path.exists():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if manifest.get("source_sha256") != current_hash:
            raise TranslatorError(
                "Word 文件在生成进度后发生变化。为避免段落错位，请备份并删除对应 progress 目录后重新 inspect。"
            )
        return manifest
    manifest = build_manifest(docx_path, config)
    atomic_write_json(manifest_path, manifest)
    return manifest


def seed_translations(manifest: dict[str, Any], path: Path) -> dict[str, Any]:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    translations: dict[str, Any] = {
        "source_sha256": manifest["source_sha256"],
        "items": {},
    }
    for item in manifest["items"]:
        static = STATIC_TRANSLATIONS.get(normalize(item["source"]))
        if static:
            translations["items"][item["id"]] = {
                "source": item["source"],
                "translation": static,
                "status": "completed",
                "method": "static",
            }
    atomic_write_json(path, translations)
    return translations


@dataclass
class LMStudioClient:
    config: dict[str, Any]

    @property
    def base_url(self) -> str:
        return str(self.config["server_url"]).rstrip("/")

    def check(self) -> None:
        try:
            response = requests.get(f"{self.base_url}/v1/models", timeout=15)
            response.raise_for_status()
            data = response.json()
        except (requests.RequestException, ValueError) as exc:
            raise TranslatorError(f"无法连接 LM Studio：{exc}") from exc
        ids = [item.get("id") for item in data.get("data", [])]
        if self.config["model"] not in ids:
            raise TranslatorError(
                f"LM Studio 未加载目标模型。需要：{self.config['model']}；当前：{ids or '无'}"
            )

    def chat(self, input_text: str, system_prompt: str, max_tokens: int | None = None) -> tuple[str, dict[str, Any]]:
        payload = {
            "model": self.config["model"],
            "input": input_text,
            "system_prompt": system_prompt,
            "temperature": self.config.get("temperature", 0.1),
            "top_p": self.config.get("top_p", 0.95),
            "top_k": self.config.get("top_k", 40),
            "min_p": self.config.get("min_p", 0.05),
            "repeat_penalty": self.config.get("repeat_penalty", 1.0),
            "max_output_tokens": max_tokens or self.config.get("max_output_tokens", 2048),
            "reasoning": self.config.get("reasoning", "off"),
            "store": False,
            "stream": False,
        }
        retries = int(self.config.get("request_retries", 3))
        timeout = int(self.config.get("request_timeout_seconds", 600))
        last_error: Exception | None = None
        for attempt in range(1, retries + 1):
            try:
                response = requests.post(
                    f"{self.base_url}/api/v1/chat",
                    json=payload,
                    timeout=timeout,
                )
                if response.status_code == 500 and "peg-native format" in response.text:
                    raise PegNativeFormatError(
                        "LM Studio 已生成文本，但 peg-native 格式解析失败。"
                    )
                response.raise_for_status()
                data = response.json()
                messages = [
                    item.get("content", "")
                    for item in data.get("output", [])
                    if item.get("type") == "message"
                ]
                content = "\n".join(part for part in messages if part).strip()
                if not content:
                    raise TranslatorError(f"模型没有返回正文：{data.get('output', [])}")
                return content, data.get("stats", {})
            except PegNativeFormatError:
                # Retrying the identical request produces the same parser error.
                # Let the translation layer split the batch instead.
                raise
            except (requests.RequestException, ValueError, TranslatorError) as exc:
                last_error = exc
                if attempt < retries:
                    time.sleep(min(2**attempt, 8))
        raise TranslatorError(f"LM Studio 请求失败（已重试 {retries} 次）：{last_error}")


def chunks_by_words(items: list[dict[str, Any]], limit: int) -> list[list[dict[str, Any]]]:
    chunks: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    current_words = 0
    current_group: str | None = None
    for item in items:
        group = "notes" if item["kind"] in {"footnote", "endnote"} else "body"
        words = max(1, int(item.get("words", word_count(item["source"]))))
        if current and (current_words + words > limit or group != current_group):
            chunks.append(current)
            current = []
            current_words = 0
        current.append(item)
        current_words += words
        current_group = group
    if current:
        chunks.append(current)
    return chunks


def load_glossary() -> list[dict[str, Any]]:
    path = GLOSSARY_DIR / "auto_glossary.csv"
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def glossary_fingerprint(rows: list[dict[str, Any]]) -> str:
    # Frequency can be recalculated after matching-rule improvements without
    # invalidating semantic audit results.
    stable_rows = [
        {
            "source": row.get("source", ""),
            "target": row.get("target", ""),
            "category": row.get("category", ""),
        }
        for row in rows
    ]
    payload = json.dumps(stable_rows, ensure_ascii=False, sort_keys=True).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def source_variant_key(value: str) -> str:
    """Normalize harmless source punctuation without collapsing short/full names."""
    value = normalize(value).replace("’", "'")
    value = re.sub(r"\bst\.\s*", "st ", value)
    value = re.sub(r"\bss\.\s*", "ss ", value)
    value = re.sub(r"[^\w\s]", " ", value)
    return " ".join(value.split())


def contexts_for_term(term: str, corpus_items: list[dict[str, Any]], maximum: int = 2) -> list[str]:
    contexts: list[str] = []
    for item in corpus_items:
        source = item.get("source", "")
        matches = term_occurrences(term, source)
        if not matches:
            continue
        match = matches[0]
        start = max(0, match.start() - 180)
        end = min(len(source), match.end() + 180)
        excerpt = source[start:end].strip()
        if start:
            excerpt = "…" + excerpt
        if end < len(source):
            excerpt += "…"
        if excerpt not in contexts:
            contexts.append(excerpt)
        if len(contexts) >= maximum:
            break
    return contexts


def related_glossary_sources(row: dict[str, Any], rows: list[dict[str, Any]]) -> list[str]:
    source = row.get("source", "")
    folded = source.casefold()
    variant = source_variant_key(source)
    category = row.get("category", "")
    related: list[str] = []
    for candidate in rows:
        other = candidate.get("source", "")
        if not other or other.casefold() == folded:
            continue
        other_folded = other.casefold()
        same_variant = source_variant_key(other) == variant
        nested = (folded in other_folded or other_folded in folded) and candidate.get("category") == category
        same_target = candidate.get("target", "") == row.get("target", "")
        if same_variant or nested or same_target:
            related.append(f"{other} => {candidate.get('target', '')}")
        if len(related) >= 6:
            break
    return related


def normalize_audit_result(result: dict[str, Any], original: dict[str, Any]) -> dict[str, Any]:
    decision = str(result.get("decision", "review")).casefold()
    if decision not in {"keep", "replace", "review"}:
        decision = "review"
    try:
        confidence = max(0.0, min(1.0, float(result.get("confidence", 0))))
    except (TypeError, ValueError):
        confidence = 0.0
    recommended = " ".join(str(result.get("recommended_target", "")).split()).strip()
    if not recommended:
        recommended = original.get("target", "")
        decision = "review"
        confidence = 0.0
    category = str(result.get("category", original.get("category", "other"))).casefold()
    allowed_categories = {
        "person", "artwork", "place", "institution", "religion", "technique", "concept", "other"
    }
    if category not in allowed_categories:
        category = original.get("category", "other")
    canonical = " ".join(str(result.get("canonical_source", original.get("source", ""))).split()).strip()
    return {
        "source": original.get("source", ""),
        "current_target": original.get("target", ""),
        "recommended_target": recommended,
        "canonical_source": canonical or original.get("source", ""),
        "current_category": original.get("category", "other"),
        "recommended_category": category,
        "frequency": original.get("frequency", "0"),
        "decision": decision,
        "confidence": f"{confidence:.2f}",
        "reason": " ".join(str(result.get("reason", "")).split()).strip(),
        "rule_flags": "",
    }


def audit_one_glossary_batch(
    client: LMStudioClient,
    batch: list[dict[str, Any]],
    all_rows: list[dict[str, Any]],
    corpus_items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    request_items = []
    for row in batch:
        request_items.append(
            {
                "source": row.get("source", ""),
                "current_target": row.get("target", ""),
                "category": row.get("category", "other"),
                "frequency": int(row.get("frequency", 0) or 0),
                "related_terms": related_glossary_sources(row, all_rows),
                "contexts": contexts_for_term(row.get("source", ""), corpus_items),
            }
        )
    raw, _ = client.chat(
        json.dumps(request_items, ensure_ascii=False, indent=2),
        GLOSSARY_AUDIT_PROMPT,
        max_tokens=2048,
    )
    results = extract_json_array(raw)
    indexed = {
        str(item.get("source", "")).casefold(): item
        for item in results
        if item.get("source")
    }
    missing = [row for row in batch if row.get("source", "").casefold() not in indexed]
    if missing and len(batch) > 1:
        # Recover only omitted entries, keeping already valid results.
        for row in missing:
            recovered = audit_one_glossary_batch(client, [row], all_rows, corpus_items)
            indexed[row.get("source", "").casefold()] = recovered[0]
    normalized: list[dict[str, Any]] = []
    for row in batch:
        value = indexed.get(row.get("source", "").casefold())
        if not value:
            value = {
                "source": row.get("source", ""),
                "recommended_target": row.get("target", ""),
                "decision": "review",
                "confidence": 0,
                "reason": "模型未返回该词条",
            }
        # Recovered single-entry calls are already normalized.
        if "current_category" in value:
            normalized.append(value)
        else:
            normalized.append(normalize_audit_result(value, row))
    return normalized


def add_audit_rule_flags(rows: list[dict[str, Any]]) -> None:
    variant_groups: dict[str, list[dict[str, Any]]] = {}
    target_groups: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        variant_groups.setdefault(source_variant_key(row["source"]), []).append(row)
        target_groups.setdefault(normalize(row["recommended_target"]), []).append(row)

    for group in variant_groups.values():
        if len(group) < 2:
            continue
        targets = {normalize(row["recommended_target"]) for row in group}
        flag = "格式变体已识别" if len(targets) == 1 else "格式变体译名冲突"
        for row in group:
            row["rule_flags"] = flag
            if len(targets) > 1:
                row["decision"] = "review"

    for group in target_groups.values():
        if len(group) < 2:
            continue
        categories = {row["recommended_category"] for row in group}
        if len(categories) > 1:
            for row in group:
                extra = "同一译名对应不同类别"
                row["rule_flags"] = ";".join(filter(None, [row["rule_flags"], extra]))
                row["decision"] = "review"

    for row in rows:
        # A known high-risk mismatch is detected mechanically, without deciding its final Chinese name.
        if "arena fresco" in row["source"].casefold() and "布兰卡切尼" in row["current_target"]:
            row["rule_flags"] = ";".join(filter(None, [row["rule_flags"], "地点疑似错配"]))
            row["decision"] = "review"


def write_csv_atomic(path: Path, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    fd, temp_name = tempfile.mkstemp(prefix=path.name, suffix=".tmp", dir=path.parent)
    try:
        with os.fdopen(fd, "w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)
        os.replace(temp_name, path)
    finally:
        if os.path.exists(temp_name):
            os.unlink(temp_name)


def select_pareto_glossary(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Select the small set most likely to affect consistency across the document."""
    top_count = max(20, math.ceil(len(rows) * 0.20))
    selected = {row["source"].casefold(): row for row in rows[:top_count]}

    variant_groups: dict[str, list[dict[str, Any]]] = {}
    target_groups: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        variant_groups.setdefault(source_variant_key(row["source"]), []).append(row)
        target_groups.setdefault(normalize(row["target"]), []).append(row)
    for group in variant_groups.values():
        if len(group) > 1:
            selected.update({row["source"].casefold(): row for row in group})
    for group in target_groups.values():
        if len(group) > 1 and len({row.get("category") for row in group}) > 1:
            selected.update({row["source"].casefold(): row for row in group})

    return sorted(
        selected.values(),
        key=lambda row: (-int(row.get("frequency", 0) or 0), row["source"].casefold()),
    )


def relevant_glossary(glossary: list[dict[str, Any]], text: str, maximum: int = 30) -> list[dict[str, Any]]:
    matches = [
        item for item in glossary
        if item.get("source", "") and term_occurrences(item.get("source", ""), text)
    ]
    matches.sort(key=lambda item: int(item.get("frequency", 0) or 0), reverse=True)
    if len(matches) < 8:
        fallback = sorted(
            glossary,
            key=lambda item: int(item.get("frequency", 0) or 0),
            reverse=True,
        )
        seen = {item.get("source", "").casefold() for item in matches}
        for item in fallback:
            key = item.get("source", "").casefold()
            if key and key not in seen:
                matches.append(item)
                seen.add(key)
            if len(matches) >= 8:
                break
    return matches[:maximum]


def format_translation_request(items: list[dict[str, Any]], glossary: list[dict[str, Any]]) -> str:
    source_text = "\n".join(item["source"] for item in items)
    selected = relevant_glossary(glossary, source_text)
    glossary_text = "（本批无已提取术语）"
    if selected:
        glossary_text = "\n".join(
            f"- {item.get('source', '')} => {item.get('target', '')}" for item in selected
        )
    blocks = "\n\n".join(f"[{item['id']}]\n{item['source']}" for item in items)
    return (
        "请翻译下面各编号文本。严格按 [编号] 换行 译文 的形式返回。\n"
        "不要输出代码块，不要添加标题。脚注中的书目信息和网址保持原样。\n\n"
        f"术语表：\n{glossary_text}\n\n"
        f"待翻译文本：\n{blocks}"
    )


def parse_translation_blocks(text: str, expected_ids: Iterable[str]) -> dict[str, str]:
    parsed = {
        match.group("id"): match.group("text").strip()
        for match in BLOCK_RE.finditer(text)
        if match.group("text").strip()
    }
    expected = list(expected_ids)
    missing = [item_id for item_id in expected if item_id not in parsed]
    if missing:
        raise TranslatorError(f"模型返回缺少编号：{', '.join(missing)}")
    return {item_id: parsed[item_id] for item_id in expected}


def translate_one_chunk(
    client: LMStudioClient,
    items: list[dict[str, Any]],
    glossary: list[dict[str, Any]],
) -> tuple[dict[str, str], dict[str, Any]]:
    request = format_translation_request(items, glossary)
    try:
        raw, stats = client.chat(request, SYSTEM_PROMPT)
    except PegNativeFormatError as exc:
        if len(items) == 1:
            raise TranslatorError(
                f"单个项目 {items[0]['id']} 仍触发 peg-native 格式错误；请保留进度并提供该次日志。"
            ) from exc
        midpoint = max(1, len(items) // 2)
        left, left_stats = translate_one_chunk(client, items[:midpoint], glossary)
        right, right_stats = translate_one_chunk(client, items[midpoint:], glossary)
        return (
            {**left, **right},
            {
                "recovered_by_batch_split": True,
                "original_batch_items": len(items),
                "left": left_stats,
                "right": right_stats,
            },
        )
    try:
        return parse_translation_blocks(raw, [item["id"] for item in items]), stats
    except TranslatorError:
        if len(items) == 1:
            raise
        # A malformed multi-block response is recovered with one request per item.
        recovered: dict[str, str] = {}
        combined_stats: dict[str, Any] = {"recovered_individually": True}
        for item in items:
            single_raw, _ = client.chat(format_translation_request([item], glossary), SYSTEM_PROMPT)
            recovered.update(parse_translation_blocks(single_raw, [item["id"]]))
        return recovered, combined_stats


def extract_json_array(text: str) -> list[dict[str, Any]]:
    cleaned = re.sub(r"^\s*```(?:json)?\s*|\s*```\s*$", "", text.strip(), flags=re.I)
    start = cleaned.find("[")
    end = cleaned.rfind("]")
    if start < 0 or end < start:
        raise TranslatorError("术语响应中没有 JSON 数组。")
    value = json.loads(cleaned[start : end + 1])
    if not isinstance(value, list):
        raise TranslatorError("术语响应不是 JSON 数组。")
    return [item for item in value if isinstance(item, dict)]


def merge_and_write_glossary(
    completed_chunks: dict[str, list[dict[str, Any]]], corpus: str
) -> int:
    merged: dict[str, dict[str, Any]] = {}
    for candidates in completed_chunks.values():
        for item in candidates:
            source = " ".join(str(item.get("source", "")).split()).strip()
            target = " ".join(str(item.get("target", "")).split()).strip()
            if not source or not target or word_count(source) == 0:
                continue
            key = source.casefold()
            if key not in merged:
                merged[key] = {
                    "source": source,
                    "target": target,
                    "category": str(item.get("category", "other")),
                }

    for item in merged.values():
        item["frequency"] = len(term_occurrences(item["source"], corpus))
    rows = sorted(merged.values(), key=lambda item: (-item["frequency"], item["source"].casefold()))
    path = GLOSSARY_DIR / "auto_glossary.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["source", "target", "category", "frequency"])
        writer.writeheader()
        writer.writerows(rows)
    return len(rows)


def command_inspect(docx_path: Path, config: dict[str, Any]) -> None:
    manifest = load_or_create_manifest(docx_path, config)
    counts = Counter(item["kind"] for item in manifest["items"])
    words = Counter()
    for item in manifest["items"]:
        words[item["kind"]] += int(item.get("words", 0))
    report = {
        "文件": manifest["source_name"],
        "SHA256": manifest["source_sha256"],
        "文档结构": manifest["document"],
        "识别边界": manifest["sections"],
        "待处理项目": dict(counts),
        "待处理英文词数（近似）": dict(words),
        "合计待处理词数（近似）": sum(words.values()),
        "说明": "参考文献条目保持原文；脚注纳入翻译；目录页码需在 Word 中最终更新。",
    }
    report_path = state_dir_for(docx_path) / "inspection_report.json"
    atomic_write_json(report_path, report)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"\n检查报告：{report_path}")


def command_terms(docx_path: Path, config: dict[str, Any], max_chunks: int | None) -> None:
    manifest = load_or_create_manifest(docx_path, config)
    client = LMStudioClient(config)
    client.check()
    sources = [
        item
        for item in manifest["items"]
        if item["kind"] in {"paragraph", "table_cell"}
        and normalize(item["source"]) not in STATIC_TRANSLATIONS
    ]
    chunks = chunks_by_words(sources, int(config.get("term_scan_words", 1800)))
    state_path = GLOSSARY_DIR / "term_scan_state.json"
    if state_path.exists():
        state = json.loads(state_path.read_text(encoding="utf-8"))
        if state.get("source_sha256") != manifest["source_sha256"]:
            raise TranslatorError("术语扫描状态属于另一版本的 Word 文件，请先备份并删除 term_scan_state.json。")
    else:
        state = {
            "source_sha256": manifest["source_sha256"],
            "total_chunks": len(chunks),
            "completed_chunks": {},
        }

    processed = 0
    for index, chunk in enumerate(chunks):
        key = str(index)
        if key in state["completed_chunks"]:
            continue
        text = "\n\n".join(item["source"] for item in chunk)
        raw, stats = client.chat(text, TERM_PROMPT, max_tokens=1200)
        candidates = extract_json_array(raw)
        state["completed_chunks"][key] = candidates
        state["last_stats"] = stats
        atomic_write_json(state_path, state)
        processed += 1
        print(f"术语扫描 {len(state['completed_chunks'])}/{len(chunks)}，本批候选 {len(candidates)} 项")
        if max_chunks is not None and processed >= max_chunks:
            break

    corpus = "\n".join(item["source"] for item in sources)
    count = merge_and_write_glossary(state["completed_chunks"], corpus)
    print(f"已生成 {GLOSSARY_DIR / 'auto_glossary.csv'}，当前合并候选 {count} 项。")
    if len(state["completed_chunks"]) < len(chunks):
        print("术语扫描尚未完成；再次运行同一命令会从断点继续。")


def command_audit_terms(
    docx_path: Path,
    config: dict[str, Any],
    max_batches: int | None,
    full: bool = False,
) -> None:
    manifest = load_or_create_manifest(docx_path, config)
    all_glossary = load_glossary()
    if not all_glossary:
        raise TranslatorError("找不到 glossary/auto_glossary.csv，请先运行 terms。")
    glossary = all_glossary if full else select_pareto_glossary(all_glossary)

    client = LMStudioClient(config)
    client.check()
    corpus_items = [item for item in manifest["items"] if item.get("source")]
    batch_size = int(config.get("term_audit_batch_size", 10))
    batches = [glossary[index : index + batch_size] for index in range(0, len(glossary), batch_size)]
    fingerprint = glossary_fingerprint(glossary)
    state_path = GLOSSARY_DIR / ("term_audit_full_state.json" if full else "term_audit_pareto_state.json")
    if state_path.exists():
        state = json.loads(state_path.read_text(encoding="utf-8"))
        if state.get("source_sha256") != manifest["source_sha256"]:
            raise TranslatorError("术语审校状态属于另一个 Word 文件。请备份后删除 term_audit_state.json。")
        if state.get("glossary_fingerprint") != fingerprint:
            raise TranslatorError("术语表在审校开始后发生变化。请备份后删除 term_audit_state.json 再重新审校。")
    else:
        state = {
            "source_sha256": manifest["source_sha256"],
            "glossary_fingerprint": fingerprint,
            "total_batches": len(batches),
            "completed_batches": {},
        }

    cached_by_source: dict[str, dict[str, Any]] = {}
    if not full:
        for cache_path in (GLOSSARY_DIR / "term_audit_state.json", GLOSSARY_DIR / "term_audit_full_state.json"):
            if not cache_path.exists():
                continue
            try:
                cache = json.loads(cache_path.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                continue
            if cache.get("source_sha256") != manifest["source_sha256"]:
                continue
            for cached_batch in cache.get("completed_batches", {}).values():
                for row in cached_batch:
                    source = str(row.get("source", "")).casefold()
                    if source:
                        cached_by_source[source] = row

    processed = 0
    for index, batch in enumerate(batches):
        key = str(index)
        if key in state["completed_batches"]:
            continue
        missing = [row for row in batch if row["source"].casefold() not in cached_by_source]
        newly_audited = (
            audit_one_glossary_batch(client, missing, all_glossary, corpus_items)
            if missing else []
        )
        combined = dict(cached_by_source)
        combined.update({row["source"].casefold(): row for row in newly_audited})
        audited = [combined[row["source"].casefold()] for row in batch]
        state["completed_batches"][key] = audited
        atomic_write_json(state_path, state)
        processed += 1
        print(f"术语审校 {len(state['completed_batches'])}/{len(batches)}，本批 {len(audited)} 项")
        if max_batches is not None and processed >= max_batches:
            break

    audited_rows: list[dict[str, Any]] = []
    for index in range(len(batches)):
        audited_rows.extend(state["completed_batches"].get(str(index), []))
    add_audit_rule_flags(audited_rows)

    fields = [
        "source", "current_target", "recommended_target", "canonical_source",
        "current_category", "recommended_category", "frequency", "decision",
        "confidence", "reason", "rule_flags",
    ]
    write_csv_atomic(GLOSSARY_DIR / "audited_glossary.csv", audited_rows, fields)

    review_rows = [
        row for row in audited_rows
        if row["decision"] != "keep"
        or normalize(row["current_target"]) != normalize(row["recommended_target"])
        or "冲突" in row["rule_flags"]
        or "错配" in row["rule_flags"]
        or float(row["confidence"] or 0) < 0.80
    ]
    write_csv_atomic(GLOSSARY_DIR / "term_review.csv", review_rows, fields)

    canonical_groups: dict[str, list[dict[str, Any]]] = {}
    for row in audited_rows:
        canonical_groups.setdefault(normalize(row["canonical_source"]), []).append(row)
    group_rows: list[dict[str, Any]] = []
    for group in canonical_groups.values():
        if len(group) < 2:
            continue
        group_rows.append(
            {
                "canonical_source": group[0]["canonical_source"],
                "source_variants": " | ".join(row["source"] for row in group),
                "current_targets": " | ".join(dict.fromkeys(row["current_target"] for row in group)),
                "recommended_targets": " | ".join(dict.fromkeys(row["recommended_target"] for row in group)),
                "needs_review": "yes" if any(row["decision"] == "review" for row in group) else "no",
            }
        )
    write_csv_atomic(
        GLOSSARY_DIR / "canonical_groups.csv",
        group_rows,
        ["canonical_source", "source_variants", "current_targets", "recommended_targets", "needs_review"],
    )

    summary = {
        "source_terms": len(all_glossary),
        "selected_terms": len(glossary),
        "mode": "full" if full else "pareto",
        "audited_terms": len(audited_rows),
        "completed_batches": len(state["completed_batches"]),
        "total_batches": len(batches),
        "keep": sum(row["decision"] == "keep" for row in audited_rows),
        "replace": sum(row["decision"] == "replace" for row in audited_rows),
        "review": sum(row["decision"] == "review" for row in audited_rows),
        "review_file_rows": len(review_rows),
        "canonical_groups": len(group_rows),
        "note": "二八审校只检查高影响和异常词条；结果尚未覆盖 auto_glossary.csv。",
    }
    atomic_write_json(GLOSSARY_DIR / "term_audit_summary.json", summary)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    print(f"完整审校：{GLOSSARY_DIR / 'audited_glossary.csv'}")
    print(f"待确认项：{GLOSSARY_DIR / 'term_review.csv'}")
    print(f"变体分组：{GLOSSARY_DIR / 'canonical_groups.csv'}")
    if len(state["completed_batches"]) < len(batches):
        print("术语审校尚未完成；再次运行同一命令会从断点继续。")


def append_log(entry: dict[str, Any]) -> None:
    path = LOG_DIR / "translation.jsonl"
    with path.open("a", encoding="utf-8", newline="\n") as handle:
        handle.write(json.dumps(entry, ensure_ascii=False) + "\n")


def command_translate(docx_path: Path, config: dict[str, Any], max_chunks: int | None) -> None:
    manifest = load_or_create_manifest(docx_path, config)
    state_dir = state_dir_for(docx_path)
    translations_path = state_dir / "translations.json"
    translations = seed_translations(manifest, translations_path)
    client = LMStudioClient(config)
    client.check()
    glossary = load_glossary()
    if not glossary:
        print("提示：尚无 auto_glossary.csv，将不带术语表翻译；建议先运行 terms。")

    pending = [
        item for item in manifest["items"] if item["id"] not in translations["items"]
    ]
    chunks = chunks_by_words(pending, int(config.get("chunk_words", 450)))
    if not chunks:
        print("所有项目均已翻译完成。")
        return

    processed = 0
    for index, chunk in enumerate(chunks, start=1):
        started = time.time()
        parsed, stats = translate_one_chunk(client, chunk, glossary)
        for item in chunk:
            translations["items"][item["id"]] = {
                "source": item["source"],
                "translation": parsed[item["id"]],
                "status": "completed",
                "method": "model",
            }
        atomic_write_json(translations_path, translations)
        elapsed = time.time() - started
        append_log(
            {
                "time": time.strftime("%Y-%m-%d %H:%M:%S"),
                "document": docx_path.name,
                "ids": [item["id"] for item in chunk],
                "elapsed_seconds": round(elapsed, 2),
                "stats": stats,
            }
        )
        processed += 1
        completed = len(translations["items"])
        total = len(manifest["items"])
        print(f"翻译批次 {index}/{len(chunks)} 完成；项目进度 {completed}/{total}；耗时 {elapsed:.1f}s")
        if max_chunks is not None and processed >= max_chunks:
            print("已按 --max-chunks 限制暂停；再次运行同一命令会继续。")
            break


def insert_translation_after(paragraph: Paragraph, text: str, style_name: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style_name
    run = new_paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return new_paragraph


def replace_visible_text(paragraph: Paragraph, text: str) -> None:
    # Preserve drawings, bookmarks and note-reference elements; clear only visible text nodes.
    for text_node in paragraph._p.xpath(".//w:t"):
        text_node.text = ""
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "微软雅黑")
    run_props.append(fonts)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.set(f"{{{XML_NS}}}space", "preserve")
    text_node.text = text
    run.append(text_node)
    insert_at = 1 if len(paragraph._p) and paragraph._p[0].tag == qn("w:pPr") else 0
    paragraph._p.insert(insert_at, run)


def ensure_translation_style(document: Document) -> str:
    name = "Translation Chinese"
    if name not in document.styles:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return name


def paragraph_for_item(document: Document, item: dict[str, Any]) -> Paragraph:
    if item["kind"] == "paragraph":
        return document.paragraphs[int(item["index"])]
    if item["kind"] == "table_cell":
        cell = document.tables[int(item["table_index"])].rows[int(item["row_index"])].cells[
            int(item["cell_index"])
        ]
        return cell.paragraphs[int(item["paragraph_index"])]
    raise TranslatorError(f"不是正文项目：{item['id']}")


def patch_note_part(
    docx_path: Path,
    part_kind: str,
    note_translations: dict[int, str],
    bilingual: bool,
) -> None:
    if not note_translations:
        return
    part_name = f"word/{part_kind}.xml"
    note_tag = "footnote" if part_kind == "footnotes" else "endnote"
    temp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        if part_name not in source_zip.namelist():
            return
        root = ET.fromstring(source_zip.read(part_name))
        nodes = {
            int(node.attrib.get(W_ID, "-1")): node
            for node in root.findall(f"w:{note_tag}", NS)
        }
        for note_id, translation in note_translations.items():
            node = nodes.get(note_id)
            if node is None:
                continue
            if bilingual:
                paragraph = ET.Element(f"{{{W_NS}}}p")
                run = ET.SubElement(paragraph, f"{{{W_NS}}}r")
                text = ET.SubElement(run, f"{{{W_NS}}}t")
                text.set(f"{{{XML_NS}}}space", "preserve")
                text.text = f"译文：{translation}"
                node.append(paragraph)
            else:
                for text_node in node.findall(".//w:t", NS):
                    text_node.text = ""
                paragraphs = node.findall("w:p", NS)
                paragraph = paragraphs[0] if paragraphs else ET.SubElement(node, f"{{{W_NS}}}p")
                run = ET.SubElement(paragraph, f"{{{W_NS}}}r")
                text = ET.SubElement(run, f"{{{W_NS}}}t")
                text.set(f"{{{XML_NS}}}space", "preserve")
                text.text = translation
        replacement = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(temp_path, "w") as target_zip:
            for info in source_zip.infolist():
                data = replacement if info.filename == part_name else source_zip.read(info.filename)
                target_zip.writestr(info, data)
    os.replace(temp_path, docx_path)


def command_render(docx_path: Path, config: dict[str, Any], allow_partial: bool) -> None:
    manifest = load_or_create_manifest(docx_path, config)
    translations_path = state_dir_for(docx_path) / "translations.json"
    if not translations_path.exists():
        raise TranslatorError("尚无翻译记录，请先运行 translate。")
    translations = json.loads(translations_path.read_text(encoding="utf-8"))["items"]
    missing = [item["id"] for item in manifest["items"] if item["id"] not in translations]
    if missing and not allow_partial:
        raise TranslatorError(
            f"还有 {len(missing)} 项未翻译。若只想预览当前结果，请添加 --allow-partial。"
        )

    regular_items = [
        item
        for item in manifest["items"]
        if item["kind"] in {"paragraph", "table_cell"} and item["id"] in translations
    ]

    bilingual_doc = Document(docx_path)
    style_name = ensure_translation_style(bilingual_doc)
    # Reverse order avoids inserted XML affecting relative placement for adjacent paragraphs.
    for item in reversed(regular_items):
        paragraph = paragraph_for_item(bilingual_doc, item)
        insert_translation_after(paragraph, translations[item["id"]]["translation"], style_name)
    bilingual_path = OUTPUT_DIR / f"{docx_path.stem}_bilingual.docx"
    bilingual_doc.save(bilingual_path)

    chinese_doc = Document(docx_path)
    for item in regular_items:
        paragraph = paragraph_for_item(chinese_doc, item)
        replace_visible_text(paragraph, translations[item["id"]]["translation"])
    chinese_path = OUTPUT_DIR / f"{docx_path.stem}_zh.docx"
    chinese_doc.save(chinese_path)

    for part_kind, prefix in (("footnotes", "FN"), ("endnotes", "EN")):
        note_items = {
            int(item["note_id"]): translations[item["id"]]["translation"]
            for item in manifest["items"]
            if item["kind"] == part_kind[:-1] and item["id"] in translations
        }
        patch_note_part(bilingual_path, part_kind, note_items, bilingual=True)
        patch_note_part(chinese_path, part_kind, note_items, bilingual=False)

    print(f"中英对照版：{bilingual_path}")
    print(f"中文阅读版：{chinese_path}")
    if missing:
        print(f"注意：这是部分预览，仍有 {len(missing)} 项未翻译。")
    print("请在 Word 中更新目录/插图目录，并重点复核脚注标记位置与专名译法。")


def command_status(docx_path: Path, config: dict[str, Any]) -> None:
    manifest = load_or_create_manifest(docx_path, config)
    translations_path = state_dir_for(docx_path) / "translations.json"
    translations = seed_translations(manifest, translations_path)["items"]
    total = len(manifest["items"])
    completed = len(translations)
    pending_words = sum(
        int(item.get("words", 0)) for item in manifest["items"] if item["id"] not in translations
    )
    print(f"项目进度：{completed}/{total}（{completed / total * 100:.1f}%）")
    print(f"剩余英文词数（近似）：{pending_words}")
    glossary = load_glossary()
    print(f"自动术语数量：{len(glossary)}")


def command_retranslate_term(
    docx_path: Path,
    config: dict[str, Any],
    source_term: str,
    target_term: str,
) -> None:
    """Update one glossary entry and invalidate only affected model translations."""
    source_term = " ".join(source_term.split()).strip()
    target_term = " ".join(target_term.split()).strip()
    if not source_term or not target_term:
        raise TranslatorError("原术语和新译名都不能为空。")

    manifest = load_or_create_manifest(docx_path, config)
    glossary_path = GLOSSARY_DIR / "auto_glossary.csv"
    rows = load_glossary()
    found = False
    for row in rows:
        if row.get("source", "").casefold() == source_term.casefold():
            row["target"] = target_term
            row["category"] = row.get("category") or "manual"
            found = True
    if not found:
        frequency = sum(len(term_occurrences(source_term, item["source"])) for item in manifest["items"])
        rows.append(
            {
                "source": source_term,
                "target": target_term,
                "category": "manual",
                "frequency": frequency,
            }
        )
    rows.sort(key=lambda row: (-int(row.get("frequency", 0) or 0), row.get("source", "").casefold()))
    with glossary_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["source", "target", "category", "frequency"])
        writer.writeheader()
        writer.writerows(rows)

    translations_path = state_dir_for(docx_path) / "translations.json"
    translations = seed_translations(manifest, translations_path)
    affected = [
        item
        for item in manifest["items"]
        if term_occurrences(source_term, item["source"])
    ]
    removed = 0
    for item in affected:
        existing = translations["items"].get(item["id"])
        if existing and existing.get("method") != "static":
            del translations["items"][item["id"]]
            removed += 1
    atomic_write_json(translations_path, translations)
    print(f"术语已更新：{source_term} => {target_term}")
    print(f"匹配原文项目：{len(affected)}；已标记重新翻译：{removed}")
    if affected:
        print("下一步运行 translate，再运行 render。原始 Word 不会被修改。")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="通过 LM Studio 本地 API 翻译 Word 论文")
    subparsers = parser.add_subparsers(dest="command", required=True)
    for name in ("inspect", "status"):
        sub = subparsers.add_parser(name)
        sub.add_argument("document", help="Word 文件路径，例如 input/document.docx")

    terms = subparsers.add_parser("terms")
    terms.add_argument("document")
    terms.add_argument("--max-chunks", type=int, help="本次最多扫描多少批，用于小范围测试")

    audit_terms = subparsers.add_parser("audit-terms")
    audit_terms.add_argument("document")
    audit_terms.add_argument("--max-batches", type=int, help="本次最多审校多少批，用于测试或分段运行")
    audit_terms.add_argument("--full", action="store_true", help="审校全部术语；默认只审校高影响的二八集合")

    translate = subparsers.add_parser("translate")
    translate.add_argument("document")
    translate.add_argument("--max-chunks", type=int, help="本次最多翻译多少批")

    render = subparsers.add_parser("render")
    render.add_argument("document")
    render.add_argument("--allow-partial", action="store_true", help="允许生成尚未完成的预览版")

    retranslate = subparsers.add_parser("retranslate-term")
    retranslate.add_argument("document")
    retranslate.add_argument("source_term", help="英文原术语")
    retranslate.add_argument("target_term", help="新的中文译名")
    return parser


def main() -> int:
    ensure_dirs()
    parser = build_parser()
    args = parser.parse_args()
    try:
        config = load_config()
        docx_path = resolve_docx(args.document)
        if args.command == "inspect":
            command_inspect(docx_path, config)
        elif args.command == "terms":
            command_terms(docx_path, config, args.max_chunks)
        elif args.command == "audit-terms":
            command_audit_terms(docx_path, config, args.max_batches, args.full)
        elif args.command == "translate":
            command_translate(docx_path, config, args.max_chunks)
        elif args.command == "render":
            command_render(docx_path, config, args.allow_partial)
        elif args.command == "status":
            command_status(docx_path, config)
        elif args.command == "retranslate-term":
            command_retranslate_term(docx_path, config, args.source_term, args.target_term)
        return 0
    except (TranslatorError, OSError, zipfile.BadZipFile, json.JSONDecodeError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
